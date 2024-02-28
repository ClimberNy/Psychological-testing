from docx import Document
from copy import deepcopy
import datetime
import re
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
#设置表格边框实线
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

#获取文件路径
def find_files_by_partial_name(directory, partial_name):
    found_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if partial_name in file:
                found_files.append(os.path.join(root, file))
    return found_files


M = find_files_by_partial_name("./source","分诊问卷")[0]
N = find_files_by_partial_name("./source","序号")[0]
K = "source/1精准心理治疗评估报告_新模板.docx"

#心理魔方
Content = {'高高高':'思考、感受、行动取向干预均可使用',
           '高高中':'优先使用思考取向干预，适度使用感受、行动取向干预',
           '高高低':'优先使用思考取向干预、适度使用行动取向干预',
           '高中高':'思考、感受、行动取向干预均可使用',
           '高中中':'优先使用思考取向干预，适度使用感受、行动取向干预',
           '高中低':'优先使用思考取向干预，适度使用行动取向干预，减少使用感受取向干预',
           '高低高':'思考、感受、行动取向干预均可使用',
           '高低中':'优先使用思考取向干预，适度使用感受、行动取向干预',
           '高低低':'优先使用思考取向干预，适度使用行动取向干预，减少使用感受取向干预',
           '中高高':'优先使用感受取向干预，适度使用思考、行动取向干预',
           '中高中':'优先使用感受、行动取向干预',
           '中高低':'优先使用思考取向干预，适度使用行动取向干预，减少使用感受取向干预',
           '中中高':'优先使用感受取向干预，适度使用思考、行动取向干预',
           '中中中':'思考、感受、行动取向干预均可使用',
           '中中低':'优先使用思考取向干预，适度使用行动取向干预，减少使用感受取向干预',
           '中低高':'思考、感受、行动取向干预均可使用',
           '中低中':'优先使用思考取向干预，适度使用感受取向干预',
           '中低低':'优先使用思考取向干预，适度使用行动取向干预，减少使用感受取向干预',
           '低高高':'适度使用感受、行为取向干预，减少使用思考取向干预',
           '低高中':'优先使用行动取向干预，适度使用感受取向干预，减少使用思考干预取向',
           '低高低':'适度使用行动取向干预，减少使用思考、感受取向干预',
           '低中高':'适度使用感受、行动取向干预，减少使用思考取向干预',
           '低中中':'优先使用行动取向干预，适度使用感受取向干预，减少使用思考取向干预',
           '低中低':'适度行动取向干预，减少使用思考、感受取向干预',
           '低低高':'适度使用感受、行动取向干预，减少使用思考取向干预',
           '低低中':'优先使用行动取向干预，适度使用感受取向干预，减少使用思考取向干预',
           '低低低':'适度使用行动取向干预，减少使用思考、感受取向干预',
           }

#设置表格内文字居中，边框实线
def center():
    # 设置表格文字居中
    i = 0
    docn  = Document(K)
    for table in docn.tables:
        if i != 0:
            table.style = 'Table Grid'
            for row in table.rows:
                for cell in row.cells:
                    # 获取单元格中的段落
                    paragraph = cell.paragraphs[0]
                    # 设置段落中文字的对齐方式为居中
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            i += 1
    
    docn.save(K)


#填充病名
def delease(str,doc_name):
    table = doc_name.tables[0]  #取第一个表格
    cell_1 = table.cell(1,0)    #取到患病名称的那一格
    paragraph = cell_1.paragraphs[0]
    paragraph.add_run(str+" ")

#用以定位病人信息
def search(str):
    doc_search = Document(N)
    paras = doc_search.paragraphs
    for num in range(len(paras)) :
        if str in paras[num].text :
            num += 1
            return paras[num].text
        
#用以返回一个表格
def table_search(str):
    doc00 = Document('source/循证心理治疗证据.docx')
    for table in doc00.tables:
        if table.cell(1,0).text == str:
            return table

#填充第一个表格
def table_0(doctor,disease):
    doc_name = Document(K)
    #添加患者姓名
    table = doc_name.tables[0]  #取第一个表格
    cell_1 = table.cell(0,0)    #取到患者姓名的那一格
    str = "1.姓名："
    patient_name = search(str)
    paragraph = cell_1.paragraphs[0]
    paragraph.add_run(patient_name)


    #添加患者姓别
    table = doc_name.tables[0]  #取第一个表格
    cell_1 = table.cell(0,1)    #取到患者姓名的那一格
    str = "3.性别："
    patient_gender = search(str)
    paragraph = cell_1.paragraphs[0]
    paragraph.add_run(patient_gender)
    

    #添加患者年龄
    table = doc_name.tables[0]  #取第一个表格
    cell_1 = table.cell(0,2)    #取到患者年龄的那一格
    str = "5.年龄"
    patient_age = search(str)
    paragraph = cell_1.paragraphs[0]
    paragraph.add_run(patient_age)
    

    #添加疾病名称
    for num in disease:
        if(num=="1"):
            delease("抑郁状态",doc_name)
        elif(num=="2"):
            delease("焦虑状态",doc_name)
        elif(num=="3"):
            delease("惊恐发作",doc_name)
        elif(num=="4"):
            delease("恐怖症",doc_name)
        elif(num=="5"):
            delease("社交焦虑",doc_name)
        elif(num=="6"):
            delease("睡眠障碍",doc_name)
        elif(num=="7"):
            delease("心境（情感）障碍",doc_name)
        elif(num=="8"):
            delease("强迫状态",doc_name)
        elif(num=="9"):
            delease("创伤后应激障碍",doc_name)
        elif(num=="10"):
            delease("神经性贪食症",doc_name)
        elif(num=="11"):
            delease("神经性厌食症",doc_name)
        elif(num=="12"):
            delease("注意力缺陷",doc_name)
            

    #添加医生姓名
    cell_1 = table.cell(1,1)    #取到医生姓名的那一格
    paragraph = cell_1.paragraphs[0]
    paragraph.add_run(doctor)
    
    #添加报告日期
    cell_1 = table.cell(1,2)
    paragraph = cell_1.paragraphs[0]
    data = datetime.date.today().strftime("%Y-%m-%d")
    paragraph.add_run(data)
    doc_name.save(K)

    
#第一部分1  可优化（使用表格索引来定位）
def part1_1(disease):
    doc00 = Document('source/药物治疗与心理治疗选择的循证依据.docx')
    doc1 = Document(K)
    paras = doc1.paragraphs
    str = "1 药物治疗与心理治疗选择的循证依据"
    str1 = "5.年龄"
    patient_age = search(str1)
    age = int(patient_age)
    for num in range(len(paras)) :
        if str == paras[num].text :
            num += 1
            for A in disease:
                if(A=="1"):                                                   
                    if(age >= 18):
                        table = doc00.tables[0]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)
                    else:
                        table = doc00.tables[1]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)  
                    doc1.save(K)
                elif(A=="2"):                    
                    if(age >= 18):
                        table = doc00.tables[2]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)
                    else:
                        table = doc00.tables[3]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)  
                    doc1.save(K)
                elif(A=="3"):
                    table = doc00.tables[4]
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                    doc1.save(K)
                elif(A=="4"):
                    if(age >= 18):
                        table = doc00.tables[5]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)
                    else:
                        table = doc00.tables[6]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)  
                    doc1.save(K)
                elif(A=="5"):
                    if(age >= 18):
                        table = doc00.tables[7]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)
                    else:
                        table = doc00.tables[8]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)  
                    doc1.save(K)
                elif(A=="6"):
                    table = doc00.tables[9]
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                    doc1.save(K)
                elif(A=="7"):
                    table = doc00.tables[10]
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                    doc1.save(K)
                elif(A=="8"):
                    table = doc00.tables[11]
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                    doc1.save(K)
                elif(A=="9"):
                    table = doc00.tables[12]
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                    doc1.save(K)
                elif(A=="10"):
                    if(age >= 18):
                        table = doc00.tables[13]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)
                    else:
                        table = doc00.tables[14]
                        new_table = deepcopy(table)
                        paragraph = paras[num].insert_paragraph_before()
                        paragraph._p.addnext(new_table._element)  
                    doc1.save(K)
                elif(A=="11"):
                    table = doc00.tables[15]
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)

    doc1.save(K)

#第一部分2
def part1_2():
    doc1 = Document(K)
    doc2 = Document(M)
    target_tables = doc1.tables
    #第一个表
    a = 0
    b = 0
    num1 = 0
    sourse_table = doc2.tables[2]
    for target_table in target_tables:
        if target_table.cell(0,0).text == "药物治疗依从性":
            for source_row, target_row in zip(sourse_table.rows, target_table.rows):
                if(a != 0): 
                    
                    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
                        if(b != 0):
                            target_cell.text = source_cell.text
                        b += 1
                    num1 = num1 + float(target_row.cells[1].text)
                a += 1
            num1 = round(num1/(a-1),2)
            target_table.cell(6,1).text = str(num1)
            if(num1 < 3.2):
                target_table.cell(6,2).text = "一般"
            else:
                target_table.cell(6,2).text = "良好"

    
    #第二个表
    a = 0
    b = 0
    num2 = 0
    sourse_table = doc2.tables[1]
    for target_table in target_tables:
        if target_table.cell(0,0).text == "心理治疗适宜性":
            for source_row, target_row in zip(sourse_table.rows, target_table.rows):
                
                if(a != 0):
                    
                    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
                        
                        if(b != 0):
                            target_cell.text = source_cell.text
                        b += 1
                    num2 = num2 + float(target_row.cells[1].text)
                a += 1
            num2 = round(num2/(a-1),2)
            target_table.cell(6,1).text = str(num2)
            if(num2 < 3.5):
                target_table.cell(6,2).text = "一般"
            else:
                target_table.cell(6,2).text = "良好"
        
    doc1.save(K)

#第一部分3 
def part1_3():
    doc1 = Document(K)
    for target_table in doc1.tables:
        if(target_table.cell(0,0).text == "项目"):
            content = search("11.您的早年成长情况？")
            target_table.cell(1,1).text = content
            content = search("10.近期遭遇事件（可多选）？")
            target_table.cell(2,1).text = content  
            content = search("7.你期待的治疗方法是？")
            target_table.cell(3,1).text = content  
            content = search("8.您拒绝接受的治疗方法是？")
            target_table.cell(4,1).text = content   
            content = search("1.你曾接受过心理咨询或治疗服务吗？")
            target_table.cell(5,1).text = content  
            content = search("9.来访主题？")
            target_table.cell(6,1).text = content 
    doc1.save(K)      

#第一部分4    #默认一个病
def part1_4(disease):
    str = "4 小结与建议"
    doc1 = Document(K)
    paras = doc1.paragraphs
    str1 = "5.年龄"
    patient_age = search(str1)
    age = int(patient_age)
    for num in range(len(paras)) :
        if str == paras[num].text :
            num += 1
            #第一行
            if len(disease) == 2:
                if age>=18:
                    paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一药物治疗或单一心理治疗，或考虑联合治疗;"
                else:
                    paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗或联合治疗，或考虑单一药物治疗;"
            else:
                if disease[0] == '1':
                    if age>=18:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一药物治疗或单一心理治疗，或考虑联合治疗;"
                    else:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗或联合治疗，或考虑单一药物治疗;"
                elif disease[0] == '2':
                    if age>=18:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一药物治疗或单一心理治疗，或考虑联合治疗;"
                    else:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗或联合治疗，或考虑单一药物治疗;"
                elif disease[0] == '3':
                    paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗或联合治疗，或考虑单一药物治疗;"
                elif disease[0] == '4':
                    if age>=18:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗，或考虑单一药物治疗;"
                    else:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗;"
                elif disease[0] == '5':
                    if age>=18:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一药物治疗或单一心理治疗，或考虑联合治疗;"
                    else:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗，或考虑联合治疗;"
                elif disease[0] == '6':
                    paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗，或考虑药物治疗，或考虑联合治疗;"
                elif disease[0] == '7':
                    paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一药物治疗或联合治疗;"
                elif disease[0] == '8':
                    paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一药物治疗或单一心理治疗，或考虑联合治疗;"
                elif disease[0] == '9':
                    paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗，或考虑单一药物治疗，或考虑联合治疗;"
                elif disease[0] == '10':
                    if age>=18:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗，或考虑单一药物治疗，或考虑联合治疗;"
                    else:
                        paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗，或考虑单一药物治疗，或考虑联合治疗;"
                else:
                    paras[num].text="（1）根据该患者的诊断及其严重程度，建议单一心理治疗，或考虑单一药物治疗，或考虑联合治疗;"
            num += 1
            #第二行
            for table in doc1.tables:
                if table.cell(0,0).text == "药物治疗依从性":
                    if float(table.cell(6,1).text) <= 3.2:
                        paras[num].text="（2）该患者药物治疗依从性较低，若采用药物治疗，建议根据各维度的情况对其进行心理健康教育，以提升其服药依从性，保证药物治疗的顺利进行；"
                    else:
                        paras[num].text="（2）该患者药物治疗依从性良好，提示药物治疗的效率和效果可能较好，建议选择药物治疗;"
            num += 1
            #第三行
            for table in doc1.tables:
                if table.cell(0,0).text == "心理治疗适宜性":
                    if float(table.cell(6,1).text) <= 3.5:
                        paras[num].text="（3）该患者心理治疗适宜性总体一般，若采用心理治疗，建议根据各维度的情况对其进行心理健康教育。\n\n\n\n"
                    else:
                        paras[num].text="（3）该患者的心理治疗适宜性总体良好，提示心理治疗的效率和效果可能较好，建议选择心理治疗。\n\n\n\n"
    doc1.save(K)      

#第二部分1
def part2_1(disease):
    str = "1循证心理治疗证据等级"
    doc1 = Document(K)
    paras = doc1.paragraphs

    for num in range(len(paras)) :
        if str == paras[num].text :
            num += 1
            for A in disease:
                if(A=="1"):
                    table = table_search("抑郁障碍")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="2"):
                    table = table_search("广泛性\n焦虑障碍")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="3"):
                    table = table_search("惊恐发作")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="4"):
                    table = table_search("特定的恐怖症")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="5"):
                    table = table_search("社交恐惧")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="6"):
                    table = table_search("睡眠障碍\n（失眠症）")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="7"):
                    table = table_search("双相情感\n障碍")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="8"):
                    table = table_search("强迫障碍")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="9"):
                    table = table_search("创伤后\n应激障碍")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="10"):
                    table = table_search("神经性\n贪食症")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
                if(A=="11"):
                    table = table_search("神经性\n厌食症")
                    new_table = deepcopy(table)
                    paragraph = paras[num].insert_paragraph_before()
                    paragraph._p.addnext(new_table._element)
    doc1.save(K)
    doc1 = Document(K)
                
#第二部分2  
def part2_2():
    doc1 = Document(K)
    doc2 = Document(M)
    # 第一个表格
    for target_table in doc1.tables:
        if(target_table.cell(0,0).text == "问题解决风格"):
            target = target_table
    for sourse_table in doc2.tables:
        if(sourse_table.cell(1,0).text == "思考型"):
            sourse = sourse_table
        #各个单元格赋值
    target.cell(1,1).text = sourse.cell(1,1).text
    if(float(target.cell(1,1).text)<2.5):
        target.cell(1,2).text = "低"
    elif(float(target.cell(1,1).text)>2.5 and float(target.cell(1,1).text)<3.5):
        target.cell(1,2).text = "中"
    else:
        target.cell(1,2).text = "高"

    target.cell(2,1).text = sourse.cell(2,1).text
    if(float(target.cell(2,1).text)<2.5):
        target.cell(2,2).text = "低"
    elif(float(target.cell(2,1).text)>2.5 and float(target.cell(2,1).text)<3.5):
        target.cell(2,2).text = "中"
    else:
        target.cell(2,2).text = "高"

    target.cell(3,1).text = sourse.cell(3,1).text
    if(float(target.cell(3,1).text)<2.5):
        target.cell(3,2).text = "低"
    elif(float(target.cell(3,1).text)>2.5 and float(target.cell(3,1).text)<3.5):
        target.cell(3,2).text = "中"
    else:
        target.cell(3,2).text = "高"
    # 第二个表格
    for target_table in doc1.tables:
        if(target_table.cell(0,0).text == "治疗风格偏好"):
            target = target_table
    for sourse_table in doc2.tables:
        if(sourse_table.cell(1,0).text == "计划性"):
            sourse = sourse_table
        #各个单元格赋值
    target.cell(1,1).text = sourse.cell(1,1).text
    if(float(target.cell(1,1).text)<2.5):
        target.cell(1,2).text = "低"
    elif(float(target.cell(1,1).text)>2.5 and float(target.cell(1,1).text)<3.5):
        target.cell(1,2).text = "中"
    else:
        target.cell(1,2).text = "高"

    target.cell(2,1).text = sourse.cell(2,1).text
    if(float(target.cell(2,1).text)<2.5):
        target.cell(2,2).text = "低"
    elif(float(target.cell(2,1).text)>2.5 and float(target.cell(2,1).text)<3.5):
        target.cell(2,2).text = "中"
    else:
        target.cell(2,2).text = "高"

    target.cell(3,1).text = sourse.cell(3,1).text
    if(float(target.cell(3,1).text)<2.5):
        target.cell(3,2).text = "低"
    elif(float(target.cell(3,1).text)>2.5 and float(target.cell(3,1).text)<3.5):
        target.cell(3,2).text = "中"
    else:
        target.cell(3,2).text = "高"
    doc1.save(K)      

#第二部分3
def part2_3(disease):
    str = "3 小结和建议"
    doc1 = Document(K)
    paras = doc1.paragraphs
    for num in range(len(paras)) :
        if str == paras[num].text :
            num += 1
            #第一句
            if len(disease) == 2:
                paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知行为疗法、认知行为疗法第三浪潮、人际关系疗法、理性情绪行为疗法、短程心理动力疗法等;"
            else:
                if disease[0] == '1':
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知行为疗法、认知行为疗法第三浪潮、人际关系疗法、理性情绪行为疗法、短程心理动力疗法等;"
                elif disease[0] == '2':
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知行为疗法;"
                elif disease[0] == '3':
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知行为疗法、行为疗法;"
                elif disease[0] == '4':  #未知 特定的恐怖症
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知行为疗法;"
                elif disease[0] == '5': 
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知行为疗法;"
                elif disease[0] == '6':
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：失眠的认知行为疗法CBT-I、认知疗法、生物反馈治疗;"
                elif disease[0] == '7':
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：心理教育、家庭聚焦疗法、认知疗法、人际与社会节律疗法;"
                elif disease[0] == '8':
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知行为疗法、暴露与反应阻止疗法、认知行为疗法第三浪潮;"
                elif disease[0] == '9':
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知疗法、延迟暴露疗法;"
                elif disease[0] == '10':
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：认知行为疗法、人际关系疗法、家庭治疗;"
                else:
                    paras[num].text = "（1）对于该患者来说，具有循证心理治疗证据的疗法有：家庭治疗、认知行为疗法、全然开放辩证行为疗法;"
            num += 1
            #第二句
            for table in doc1.tables:
                if table.cell(0,0).text == "问题解决风格":
                    str = table.cell(1,2).text+table.cell(2,2).text+table.cell(3,2).text
                    paras[num].text = "（2）在心理干预取向方面，建议"+Content[str]+"；"
            num += 1
            #第三句
            for table in doc1.tables:
                if table.cell(0,0).text == "治疗风格偏好":
                    if table.cell(1,2).text == "高":
                        str3 = "较高"
                    else:
                        str3 = "适当"
                    if table.cell(2,2).text == "高":
                        str4 = "及时"
                    else:
                        str4 = "适时"
                    if table.cell(3,2).text == "高":
                        str1 = "较近"
                        str2 = "及时"
                    else:
                        str1 = "适当"
                        str2 = "适时"
                    paras[num].text = "（3）在治疗风格偏好方面，建议治疗师保持"+str1+"心理距离，"+str2+"共情，并保持"+str3+"的计划性、结构性，"+str4+"指导；"
            num += 1
            #第四句
            words = paras[num-3].text.split("：")
            words = words[1].split("、",1)
            if len(words) == 1:
                paras[num].text = "综上，建议优先使用"+words[0][0:-1]+"。"
            else:
                paras[num].text = "综上，建议优先使用"+words[0]+",也可以使用"+words[1][0:-1]+"。"
    doc1.save(K)      

#第二部分4(设置日期)
def timer():
    data = datetime.date.today().strftime("%Y-%m-%d")
    doc0 = Document(K)
    for para in doc0.paragraphs:
        if para.text == "报告日期：":
            para.add_run(data)
    doc0.save(K)




def main():
    doctor =  input("请输入医生姓名：")
    disease0 = input("1.抑郁状态\n2.焦虑状态\n3.惊恐发作\n4.恐怖症\n5.社交焦虑\
                    \n6.睡眠障碍\n7.心境（情感）障碍\n8.强迫状态\n9.创伤后应激障碍\
                    \n10.神经性贪食症\n11.神经性厌食症\n请输入疾病序号:(多个病症请用逗号隔开)\n")
    disease = re.split('，|,',disease0)
    table_0(doctor,disease)
    part1_1(disease)
    part1_2()
    part1_3()
    part1_4(disease)
    part2_1(disease)
    part2_2()
    part2_3(disease)
    center()
    timer()



main()